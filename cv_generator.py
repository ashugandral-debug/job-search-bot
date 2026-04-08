"""
CV Generator — pure Python using python-docx (no Node.js needed)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ACCENT    = RGBColor(0x1F, 0x4E, 0x79)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xEB, 0xF3, 0xFB)
GRAY      = RGBColor(0x59, 0x59, 0x59)
DARK      = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT2   = RGBColor(0xBD, 0xD7, 0xEE)


def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_thin_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_run(para, text, bold=False, italic=False, size=10,
             color=None, font="Arial"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _section_header(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    # Bottom border
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), "1F4E79")
    pBdr.append(bot)
    pPr.append(pBdr)
    _add_run(para, text.upper(), bold=True, size=11, color=ACCENT)
    return para


def generate_cv(job, profile_data):
    company   = job["company"]
    role      = job["role"]
    reqs      = job.get("key_requirements", "")
    safe_name = company.replace(" ", "_").replace(".", "").replace(",", "")
    filename  = "CV_AjatyaGandral_" + safe_name + ".docx"
    filepath  = os.path.join("cvs", filename)

    summary = _tailor_summary(profile_data["profile"], role, company, reqs)
    skills  = _select_skills(profile_data["skills"], reqs)

    try:
        _build_cv(profile_data, summary, skills, filepath)
        print("  CV generated: " + filename)
        return filepath
    except Exception as e:
        print("  CV generation failed: " + str(e))
        return None


def _tailor_summary(base_profile, role, company, requirements):
    reqs_lower = requirements.lower()
    additions  = []
    if "claims" in reqs_lower or "fidic" in reqs_lower:
        additions.append("experienced in FIDIC contract management and claims-related documentation")
    if "site" in reqs_lower or "supervision" in reqs_lower:
        additions.append("with proven site supervision and HSE compliance track record")
    if "consultancy" in reqs_lower or "consultant" in reqs_lower:
        additions.append("bringing a consultancy-oriented mindset and multi-level reporting capabilities")
    if "building" in reqs_lower or "construction" in reqs_lower:
        additions.append("specialising in building construction scheduling and milestone control")
    if "recovery" in reqs_lower or "delay" in reqs_lower:
        additions.append("known for proactive delay identification and programme recovery")
    summary = base_profile.strip()
    if additions:
        summary += " " + "; ".join(additions) + "."
    return summary


def _select_skills(all_skills, requirements):
    reqs_lower = requirements.lower()
    priority, others = [], []
    keywords = ["primavera","p6","ms project","claims","fidic","delay",
                "recovery","bim","revit","autocad","excel","site","hse"]
    for skill in all_skills:
        if any(k in skill.lower() for k in keywords if k in reqs_lower):
            priority.append(skill)
        else:
            others.append(skill)
    return (priority + others)[:18]


def _build_cv(profile_data, summary, skills, filepath):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin   = Cm(1.9)
        section.right_margin  = Cm(1.9)

    name        = profile_data["name"]
    email       = profile_data["email"]
    phone       = profile_data["phone"]
    experience  = profile_data["experience"]
    education   = profile_data["education"]
    credentials = profile_data["credentials"]

    # ── Header name bar ──────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "1F4E79")
    _set_cell_margins(cell, 160, 120, 200, 200)
    _remove_cell_borders(cell)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p1, name, bold=True, size=22, color=WHITE)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, "PLANNING ENGINEER", size=13, color=ACCENT2)

    doc.add_paragraph()

    # ── Contact bar ───────────────────────────────────────────
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.style = "Table Grid"
    cell2 = tbl2.cell(0, 0)
    _set_cell_bg(cell2, "EBF3FB")
    _set_cell_margins(cell2, 80, 80, 200, 200)
    _remove_cell_borders(cell2)

    pc = cell2.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = "Email: " + email + "   |   Phone: " + phone + "   |   UAE Attested Degree   |   Valid Dubai Driving License"
    _add_run(pc, contact, size=9, color=ACCENT)

    doc.add_paragraph()

    # ── Professional Summary ──────────────────────────────────
    _section_header(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after  = Pt(3)
    _add_run(sp, summary, size=10, color=DARK)

    # ── Core Competencies ─────────────────────────────────────
    _section_header(doc, "Core Competencies")
    cols  = 3
    rows  = (len(skills) + cols - 1) // cols
    stbl  = doc.add_table(rows=rows, cols=cols)
    stbl.style = "Table Grid"
    for i in range(rows):
        for j in range(cols):
            idx  = i * cols + j
            cell = stbl.cell(i, j)
            txt  = ("* " + skills[idx]) if idx < len(skills) else ""
            _set_cell_bg(cell, "EBF3FB")
            _set_cell_margins(cell, 60, 60, 100, 100)
            _set_thin_border(cell)
            p = cell.paragraphs[0]
            _add_run(p, txt, size=9, color=DARK)

    # ── Professional Experience ───────────────────────────────
    _section_header(doc, "Professional Experience")
    for exp in experience:
        # Title row
        etbl = doc.add_table(rows=1, cols=2)
        etbl.style = "Table Grid"
        lc = etbl.cell(0, 0)
        rc = etbl.cell(0, 1)
        _remove_cell_borders(lc)
        _remove_cell_borders(rc)

        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        _add_run(lp, exp["title"] + " | " + exp["company"], bold=True, size=10, color=DARK)

        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(4)
        _add_run(rp, exp["period"], italic=True, size=9, color=GRAY)

        loc_p = doc.add_paragraph()
        loc_p.paragraph_format.space_before = Pt(0)
        loc_p.paragraph_format.space_after  = Pt(2)
        _add_run(loc_p, exp["location"], italic=True, size=9, color=GRAY)

        for b in exp["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after  = Pt(2)
            bp.paragraph_format.left_indent  = Inches(0.25)
            _add_run(bp, b, size=9.5, color=DARK)

        doc.add_paragraph()

    # ── Education ─────────────────────────────────────────────
    _section_header(doc, "Education")
    for edu in education:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(4)
        dp.paragraph_format.space_after  = Pt(1)
        _add_run(dp, edu["degree"], bold=True, size=10, color=DARK)

        ip = doc.add_paragraph()
        ip.paragraph_format.space_before = Pt(0)
        ip.paragraph_format.space_after  = Pt(1)
        _add_run(ip, edu["institution"] + "  |  " + edu["period"], italic=True, size=9, color=GRAY)

        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after  = Pt(6)
        _add_run(fp, edu["focus"], size=9, color=GRAY)

    # ── Credentials ───────────────────────────────────────────
    _section_header(doc, "Certifications & UAE Credentials")
    for c in credentials:
        cp = doc.add_paragraph(style="List Bullet")
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(2)
        cp.paragraph_format.left_indent  = Inches(0.25)
        _add_run(cp, c, size=9.5, color=DARK)

    # ── Languages ─────────────────────────────────────────────
    _section_header(doc, "Languages")
    lp = doc.add_paragraph()
    _add_run(lp, "English — Professional Working Proficiency (Written & Spoken)   |   Hindi — Native",
             size=10, color=DARK)

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
